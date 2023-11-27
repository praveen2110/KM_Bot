from flask import Flask, render_template, request, jsonify
from docx import Document
import openai
import nltk
import json
import os
from nltk.tokenize import sent_tokenize
from fuzzywuzzy import fuzz
import re


nltk.download("punkt")
app = Flask(__name__)
valid_words = set(word.lower() for word in nltk.corpus.words.words())
# Initialize OpenAI API with your API key
#openai.api_key = "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"  # Replace with your OpenAI API key

def generate_openai_response(prompt):
    # Call OpenAI API to get a response
    openai_response = openai.Completion.create(
        engine="text-davinci-003",
        prompt=prompt,
        max_tokens=200
    )
    return openai_response.choices[0].text.strip()

def normalize_text(text):
    # Normalize text by converting to lowercase and removing non-alphanumeric characters
    return re.sub(r'[^a-zA-Z0-9\s]', '', text.lower())

def search_similar_question(user_question):
    

    fake_database = [] 

    # Normalize user's question
    normalized_user_question = normalize_text(user_question)

    # Compare normalized user's question with normalized existing questions using fuzzy matching
    similarity_threshold = 80  # Adjust the threshold as needed
    for item in fake_database:
        knowledge_info = item['title']        
        normalized_knowledge_info = normalize_text(knowledge_info)
        similarity = fuzz.token_set_ratio(normalized_user_question, normalized_knowledge_info)

        if similarity >= similarity_threshold:           
            return item['content']

    return None  

def has_tables(docx_path):
    try:
        doc = Document(docx_path)
        return any(table for table in doc.tables)
    except Exception as e:
        print(f"Error checking for tables: {e}")
        return False

def has_paragraphs(docx_path):
    try:
        doc = Document(docx_path)
        return any(paragraph for paragraph in doc.paragraphs)
    except Exception as e:
        print(f"Error checking for paragraphs: {e}")
        return False

def search_in_folder(user_question, folder_path):
    threshold = 80

    try:
        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)

            if filename.endswith(".docx"):
                if has_tables(file_path):
                    doc = Document(file_path)
                    for table in doc.tables:
                        for row in table.rows:
                            question_cell = row.cells[1]  # Assuming the question is in the first column
                            answer_cell = row.cells[4]    # Assuming the answer is in the second column

                            # Check if the user's question is similar to the question in the cell
                            if fuzz.token_set_ratio(user_question.lower(), question_cell.text.lower()) >= threshold:
                                return answer_cell.text

                elif has_paragraphs(file_path):
                    doc = Document(file_path)
                    for paragraph in doc.paragraphs:
                        # Check if the user's question is similar to the paragraph text
                        if fuzz.token_set_ratio(user_question.lower(), paragraph.text.lower()) >= threshold:
                            return paragraph.text

    except Exception as e:
        print(f"Error reading DOCX files in the folder: {e}")

    return None

def check_valid_word(message):
    tokens = nltk.word_tokenize(message)
    is_valid_word = any(token.lower() in valid_words for token in tokens)
    return is_valid_word

def extract_relevant_information(openai_text, user_message):
    # Tokenize the OpenAI response into sentences
    sentences = sent_tokenize(openai_text)

    # Tokenize the user's message into words
    user_tokens = nltk.word_tokenize(user_message.lower())

    # Find the most similar sentence to the user's message
    most_similar_sentence = find_most_similar_sentence(user_tokens, sentences)

    return most_similar_sentence

def find_most_similar_sentence(user_tokens, sentences):
    # Calculate Jaccard similarity between user tokens and each sentence
    similarities = [calculate_jaccard_similarity(user_tokens, nltk.word_tokenize(sentence.lower())) for sentence in sentences]

    # Get the index of the most similar sentence
    most_similar_index = similarities.index(max(similarities))

    # Return the most similar sentence
    return sentences[most_similar_index]

def calculate_jaccard_similarity(set1, set2):
    # Calculate Jaccard similarity between two sets
    intersection = len(set(set1).intersection(set2))
    union = len(set(set1).union(set2))
    
    return intersection / union if union != 0 else 0.0

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/get_answer", methods=["POST"])
def get_answer():
    user_question = request.form["user_question"]
    user_question = normalize_text(user_question.lower())

    is_valid_word = check_valid_word(user_question)
    if is_valid_word:
        #similar_answer = search_similar_question(user_question)
        similar_answer_docx = search_in_folder(user_question,folder_path="C:\Code_Repo\KM_Bot\Database")
        # if similar_answer:
        #     chatbot_answer = similar_answer
        if similar_answer_docx:
            chatbot_answer = similar_answer_docx
        else:
            chatbot_answer = "Sorry, I couldn't find information related to "+user_question+".Is there something I can help you with?"
            # #Call OpenAI API to get a response
            # OpenAI_answer = generate_openai_response(user_question)              
            # chatbot_answer = extract_relevant_information(OpenAI_answer,user_question)
            # #chatbot_answer = filtered_response.strip() if filtered_response else "Sorry, I could not understand"               
            # is_valid_word = check_valid_word(chatbot_answer)
            # if is_valid_word:
            #     chatbot_answer = chatbot_answer
            # else:
            #     chatbot_answer = "I'm Sorry, I don't understand what you mean by "+user_question+" Is there something I can help you with?"
            
    else:
        chatbot_answer = "I'm Sorry, I don't understand what you mean by "+user_question+".Is there something I can help you with?"

    return jsonify({"answer": chatbot_answer})
               


if __name__ == "__main__":
    app.run(debug=True)
